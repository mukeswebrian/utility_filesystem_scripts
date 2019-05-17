from docx import Document
import pandas as pd

# set inputs
document = Document("IR_Template.docx")
normal_style = document.paragraphs[12].style
preamble = document.paragraphs[13].style
questions = document.paragraphs[16].style
question_style = document.paragraphs[17].style

def clean_ref(text):
    tokens = text.split(",")
    new_text = ''
    for token in tokens:
        temp = token.strip()
        if "Schedule" in temp:
            temp = temp + "-Transmission System Plan"
            
        new_text = new_text + temp + ", "
    return new_text

def write_line(line, irs, outdoc):
    
    ref = "Ref:\t"+clean_ref(irs.Reference.iloc[line])
    TBD_text = irs.Preamble.iloc[line]
    comment = irs.Questions.iloc[line]
    intro = "At the above noted reference, HONI stated the following:"
    
    outdoc.add_paragraph("2-Staff-"+str(line+1), style=preamble)
    outdoc.add_paragraph(ref, style=normal_style)
    outdoc.add_paragraph("Preamble:", style=preamble)
    outdoc.add_paragraph(intro, style=normal_style)
    outdoc.add_paragraph(TBD_text, style=normal_style)
    outdoc.add_paragraph("Questions:", style=preamble)
    outdoc.add_paragraph(comment, style=question_style)

if __name__ == "__main__":
    
    outdoc = Document("populated.docx")
    irs = pd.read_excel("comments-referenced.xlsx", sheet_name="Sheet1")

    for line in range(0, len(irs)):
        write_line(line, irs, outdoc)
    
    outdoc.save("populated1.docx")
